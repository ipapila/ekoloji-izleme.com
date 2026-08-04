#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
sosyal_medya_gonder.py — Günlük raporu sosyal medya paylaşım formatında
bir Word belgesine dönüştürür ve e-posta ile gönderir.

Format:
  [ekoloji-izleme.com logosu]
  **Başlık** (kalın)
  *X Haber + Y İhlal* (italik)
  (boş satır)
  Rapor metni (giriş + yorum + bakış)
  (boş satır)
  *Ayrıntılar ve daha fazlası için: https://ekoloji-izleme.com/* (italik, en altta)

rapor_uret.py tarafından üretilen rapor.json dosyasını girdi olarak kullanır.
rapor.yml workflow'unda rapor_uret.py'dan hemen sonra çalıştırılır.
"""

import json
import os
import smtplib
import ssl
import sys
from datetime import datetime
from email.message import EmailMessage
from io import BytesIO
from pathlib import Path
from zoneinfo import ZoneInfo

import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

RAPOR_YEREL = Path("rapor.json")
LOGO_YEREL  = Path("ekoloji-izleme_amblem_transparan.png")
LOGO_URL    = "https://ekoloji-izleme.com/ekoloji-izleme_amblem_transparan.png"
SITE_URL    = "https://ekoloji-izleme.com/"
TR_SAAT     = ZoneInfo("Europe/Istanbul")

SMTP_HOST = os.environ.get("SMTP_HOST", "smtp.gmail.com")
SMTP_PORT = int(os.environ.get("SMTP_PORT", "465"))
SMTP_USER = os.environ.get("SMTP_USER", "")
SMTP_PASS = os.environ.get("SMTP_PASS", "")
ALICI     = os.environ.get("RAPOR_ALICI", "ipapila@gmail.com")


def rapor_oku():
    if not RAPOR_YEREL.exists():
        print("HATA: rapor.json bulunamadı, e-posta gönderilmeyecek.")
        sys.exit(1)
    return json.loads(RAPOR_YEREL.read_text(encoding="utf-8"))


def logo_indir():
    # 1) Repoda checkout edilmiş yerel dosya varsa onu kullan (ağ bağımlılığı yok)
    if LOGO_YEREL.exists():
        try:
            return BytesIO(LOGO_YEREL.read_bytes())
        except Exception as e:
            print(f"⚠ Yerel logo okunamadı, siteden indirilecek: {e}")
    # 2) Yoksa canlı siteden indir
    try:
        r = requests.get(LOGO_URL, timeout=20)
        r.raise_for_status()
        return BytesIO(r.content)
    except Exception as e:
        print(f"⚠ Logo indirilemedi, logosuz devam edilecek: {e}")
        return None


def docx_olustur(rapor, tarih_str):
    vo = rapor.get("veri_ozet", {}) or {}
    haber_sayisi = vo.get("haber_sayisi", 0)
    ihlal_sayisi = vo.get("ihlal_sayisi", 0)
    baslik = rapor.get("baslik") or "Ekoloji İzleme Günlük Rapor"

    govde = "\n\n".join(
        p for p in [rapor.get("giris", ""), rapor.get("yorum", ""), rapor.get("bakia", "")] if p
    )

    doc = Document()

    # Logo — en üstte, ortalanmış
    logo = logo_indir()
    if logo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo, width=Inches(1.4))

    # Başlık — kalın
    p = doc.add_paragraph()
    run = p.add_run(f"{baslik} — {tarih_str}")
    run.bold = True
    run.font.size = Pt(14)

    # Alt satır — italik "X Haber + Y İhlal"
    p = doc.add_paragraph()
    run = p.add_run(f"{haber_sayisi} Haber + {ihlal_sayisi} İhlal")
    run.italic = True
    run.font.size = Pt(11)

    # Boş satır
    doc.add_paragraph()

    # Rapor metni
    for paragraf in govde.split("\n\n"):
        if paragraf.strip():
            p = doc.add_paragraph()
            run = p.add_run(paragraf.strip())
            run.font.size = Pt(11)

    # Boş satır
    doc.add_paragraph()

    # Alt satır — italik link
    p = doc.add_paragraph()
    run = p.add_run(f"Ayrıntılar ve daha fazlası için: {SITE_URL}")
    run.italic = True
    run.font.size = Pt(11)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf, haber_sayisi, ihlal_sayisi, baslik


def eposta_gonder(dosya_buf, dosya_adi, konu, govde_metni):
    if not SMTP_USER or not SMTP_PASS:
        print("HATA: SMTP_USER / SMTP_PASS tanımlı değil, e-posta gönderilemedi.")
        sys.exit(1)

    msg = EmailMessage()
    msg["Subject"] = konu
    msg["From"] = SMTP_USER
    msg["To"] = ALICI
    msg.set_content(govde_metni)
    msg.add_attachment(
        dosya_buf.read(),
        maintype="application",
        subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=dosya_adi,
    )

    context = ssl.create_default_context()
    with smtplib.SMTP_SSL(SMTP_HOST, SMTP_PORT, context=context) as server:
        server.login(SMTP_USER, SMTP_PASS)
        server.send_message(msg)
    print(f"✓ E-posta gönderildi: {ALICI}")


if __name__ == "__main__":
    print("=== Sosyal Medya Formatı Üretimi ve E-posta Gönderimi ===")
    rapor = rapor_oku()

    if rapor.get("hata"):
        print("Rapor hata içeriyor (boş rapor), e-posta yine de bilgilendirme amaçlı gönderilecek.")

    tarih_str = datetime.now(TR_SAAT).strftime("%d %B %Y")
    dosya_buf, haber_sayisi, ihlal_sayisi, baslik = docx_olustur(rapor, tarih_str)
    dosya_adi = f"gunluk_rapor_{datetime.now(TR_SAAT).strftime('%Y-%m-%d')}.docx"

    konu = f"✅ Günlük Rapor Hazır — {haber_sayisi} Haber + {ihlal_sayisi} İhlal ({tarih_str})"
    govde_metni = (
        f"Bugünün raporu hazır: {baslik}\n\n"
        f"{haber_sayisi} Haber + {ihlal_sayisi} İhlal\n\n"
        "Sosyal medya paylaşımına hazır format ekte Word belgesi olarak yer alıyor.\n\n"
        f"Ayrıntılar ve daha fazlası için: {SITE_URL}"
    )

    eposta_gonder(dosya_buf, dosya_adi, konu, govde_metni)
    print("=== Tamamlandı ===")
